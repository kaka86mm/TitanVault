"""
tools.py — Fast Research 工具集 (SearXNG 搜索 + 网页阅读)

本地化实现, 无外部付费 API。
- search: SearXNG 元搜索引擎
- visit: requests + trafilatura 网页抓取, 失败 fallback Jina Reader (agent-reach)
"""
import os
import re
import json
import requests
from typing import List, Optional
from urllib.parse import urlparse

SEARXNG_URL = os.environ.get("SEARXNG_URL", "http://host-gateway:8087")
MINERU_URL = os.environ.get("MINERU_URL", "http://host-gateway:18080")
# Jina Reader (agent-reach 的 web 渠道后端): 能读 JS 渲染页面, 返回干净 markdown。
# 走 mihomo 代理 (r.jina.ai 国内不可达)。
JINA_READER_URL = os.environ.get("JINA_READER_URL", "https://r.jina.ai")
# Exa 语义搜索 (agent-reach search 渠道, 通过宿主 mcporter 调用)
EXA_MCPORTER_URL = os.environ.get("EXA_MCPORTER_URL", "http://host-gateway:18061")
# agent-reach bridge (宿主 :18061, 转发到 Exa/Twitter/小红书)
REACH_BRIDGE_URL = os.environ.get("REACH_BRIDGE_URL", "http://host-gateway:18061")
# HTTP/HTTPS 代理 (mihomo, host 模式监听宿主 7890)
HTTP_PROXY = os.environ.get("HTTP_PROXY") or os.environ.get("http_proxy") or "http://host-gateway:7890"
PROXIES = {"http": HTTP_PROXY, "https": HTTP_PROXY} if HTTP_PROXY else None
# 已知需要 JS 渲染的域名 — 直接走 Jina Reader (trafilatura 抓不到)
JS_HEAVY_DOMAINS = {"twitter.com", "x.com", "reddit.com", "instagram.com",
                    "medium.com", "zhihu.com", "bilibili.com"}

# 验证用的停用词 (构造 verify query 时去掉)
STOP_WORDS = frozenset(
    "the a an is are was were be been being have has had do does did "
    "will would could should may might must can of to in on at for with "
    "and or not but if then that this these those it its as from by "
    "的 了 是 在 和 与 或 也 就 都 还 不 没 有 对 从 被 把 让 "
    "what which who when where why how about into than".split()
)


def search(query: str, skip_queries: list = None) -> dict:
    """SearXNG 搜索。返回 {"query", "results": [{"title","url","snippet"}]}.

    skip_queries: 已搜过的 query 列表 (去重, 避免迭代时重复搜索)。
    """
    # 提取搜索关键词 (长句→关键词, 提升搜索质量)
    query = _refine_query(query)

    if skip_queries and query.lower() in [q.lower() for q in skip_queries]:
        return {"query": query, "results": [], "skipped": True}

    lang = "zh" if any("\u4E00" <= c <= "\u9FFF" for c in query) else "en"
    try:
        resp = requests.get(
            f"{SEARXNG_URL}/search",
            params={"q": query, "format": "json", "categories": "general",
                    "pageno": 1, "language": lang},
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        return {"query": query, "results": [], "error": str(e)}

    results = []
    for r in data.get("results", [])[:10]:
        results.append({
            "title": r.get("title", ""),
            "url": r.get("url", ""),
            "snippet": r.get("content", "")[:300],
        })
    return {"query": query, "results": results}


def _refine_query(query: str) -> str:
    """将长句/口语化问题提炼为搜索引擎友好的关键词。

    SearXNG/Bing/Google 对长句搜索效果差, 需要提取核心关键词。
    策略:
    1. 短关键词 (<30字且无逗号) 直接用
    2. 长句: 去口语前缀 → 按标点分句取最短有效段 → 截断
    3. 纯指令性文本 (无实体名词) → 返回原值让搜索引擎处理
    """
    query = query.strip()
    # 短关键词直接用
    if len(query) <= 30 and not re.search(r'[，。！？、；：\?\!\.]', query):
        return query

    # 去掉口语前缀 (递归去, 有的多层前缀)
    for _ in range(3):
        old = query
        query = re.sub(r'^(帮我|请|你|给我|我想|我要|看看|搜一下|搜索|查一下|研究一下|调研|分析一下|你帮我|整体研究一下|补充|了解一下|麻烦)\s*', '', query)
        if query == old:
            break

    # 去掉指令性后缀
    query = re.sub(r'(引入社交媒体信息|加入.*判断|可以引入.*信息|补充.*信息|引入.*信息)$', '', query)
    # 去掉 "最新 2026年7月" 之类的后缀 (预热搜索加的)
    query = re.sub(r'\s*最新\s*\d{4}年\d+月\s*$', '', query)

    # 如果有逗号/问号, 提取实体最密集的分句
    parts = re.split(r'[，。！？\?\!；;]', query)
    parts = [p.strip() for p in parts if len(p.strip()) > 3]
    if len(parts) > 1:
        # 选包含数字、英文、专有名词最多的分句 (实体密度最高)
        def entity_score(s):
            score = 0
            score += len(re.findall(r'\d', s)) * 2      # 数字
            score += len(re.findall(r'[A-Z]', s)) * 2    # 大写英文 (专有名词)
            score += len(re.findall(r'[\u4e00-\u9fff]{2,}', s))  # 中文词组
            return score
        best = max(parts, key=entity_score)
        if len(best) <= 50:
            query = best
        else:
            query = parts[0]

    # 最终截断到 60 字符
    if len(query) > 60:
        query = query[:60]

    query = re.sub(r'\s+', ' ', query).strip()
    return query if query else query  # 空就返回空, 不编造


def wechat_search(query: str) -> dict:
    """微信公众号文章搜索 (weixin_search_mcp via bridge)。"""
    query = _refine_query(query)
    try:
        resp = requests.post(
            f"{REACH_BRIDGE_URL}/wechat",
            json={"query": query},
            timeout=40, proxies=None,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        return {"query": query, "results": [], "error": f"WeChat bridge: {e}"}

    if not data.get("ok"):
        return {"query": query, "results": [],
                "error": data.get("error", "weixin_search failed")}

    # bridge 返回 mcporter JSON: {"result": [{title, real_url, publish_time}, ...]}
    import json as _json
    try:
        parsed = _json.loads(data.get("raw", ""))
        items = parsed.get("result", [])
    except (_json.JSONDecodeError, TypeError):
        items = []

    results = []
    for item in items[:10]:
        title = item.get("title", "")
        url = item.get("real_url", "") or item.get("link", "")
        results.append({
            "title": title[:120],
            "url": url,
            "snippet": title[:200],  # 微信搜索无摘要, 用标题
        })
    return {"query": query, "results": results, "source": "wechat"}


def exa_search(query: str, num: int = 5) -> dict:
    """Exa 语义搜索 (agent-reach)。擅长英文/技术/代码, 返回高质量结果+高亮。
    Exa 是语义搜索, 不需要提炼关键词, 但仍然去掉口语前缀。
    """
    query = _refine_query(query)
    try:
        resp = requests.post(
            f"{REACH_BRIDGE_URL}/exa",
            json={"query": query, "num": num},
            timeout=50, proxies=None,  # bridge 在宿主, 不走代理
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        return {"query": query, "results": [], "error": f"Exa bridge: {e}"}

    if not data.get("ok"):
        return {"query": query, "results": [],
                "error": data.get("error", "Exa failed")}

    # bridge 返回 mcporter 的纯文本输出 (Title:/URL:/Highlights: 格式)
    raw = data.get("raw", "")
    results = _parse_exa_output(raw)
    return {"query": query, "results": results, "source": "exa"}


def _parse_exa_output(raw: str) -> list:
    """解析 mcporter exa 输出 (Title:/URL:/Highlights: 块)。"""
    results = []
    blocks = re.split(r"\n(?=Title:)", raw.strip())
    for b in blocks:
        title = re.search(r"^Title:\s*(.*)", b, re.MULTILINE)
        url = re.search(r"^URL:\s*(.*)", b, re.MULTILINE)
        highlights = re.search(r"^Highlights:\s*\n?(.*?)(?=\n[A-Z]|\Z)", b, re.DOTALL)
        if title and url:
            results.append({
                "title": title.group(1).strip()[:120],
                "url": url.group(1).strip(),
                "snippet": (highlights.group(1).strip() if highlights else "")[:200],
            })
    return results


def twitter_search(query: str, num: int = 10) -> dict:
    """Twitter/X 搜索 (agent-reach)。通过宿主 twitter-cli。
    返回 {"query", "results": [{"title","url","snippet"}]}.
    """
    query = _refine_query(query)
    try:
        resp = requests.post(
            f"{REACH_BRIDGE_URL}/twitter",
            json={"query": query, "num": num},
            timeout=50, proxies=None,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        return {"query": query, "results": [], "error": f"Twitter bridge: {e}"}

    if not data.get("ok"):
        return {"query": query, "results": [],
                "error": data.get("error", "twitter-cli failed")}

    raw = data.get("raw", "")
    results = _parse_twitter_output(raw)
    return {"query": query, "results": results, "source": "twitter"}


def _parse_twitter_output(raw: str) -> list:
    """解析 twitter-cli yaml 输出, 提取推文。"""
    results = []
    # twitter-cli yaml 每条推文有 text/url/id 字段
    # 简单按双换行分块提取
    for block in raw.split("\n---\n"):
        text_m = re.search(r"(?:text|content|full_text):\s*[\"']?(.*?)(?:[\"']?\s*$)", block, re.MULTILINE)
        url_m = re.search(r"url:\s*(https?://\S+)", block)
        if text_m:
            text = text_m.group(1).strip()
            results.append({
                "title": text[:100],
                "url": url_m.group(1).strip() if url_m else "",
                "snippet": text[:200],
            })
    return results[:10]


def xiaohongshu_search(query: str) -> dict:
    """小红书搜索 (agent-reach)。通过宿主 xiaohongshu-mcp。
    返回 {"query", "results": [{"title","url","snippet"}]}.
    """
    try:
        resp = requests.post(
            f"{REACH_BRIDGE_URL}/xiaohongshu",
            json={"query": query},
            timeout=130, proxies=None,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        return {"query": query, "results": [], "error": f"小红书 bridge: {e}"}

    if not data.get("ok"):
        return {"query": query, "results": [],
                "error": data.get("error", "xiaohongshu-mcp failed")}

    raw = data.get("raw", "")
    results = _parse_xhs_output(raw)
    return {"query": query, "results": results, "source": "xiaohongshu"}


def _parse_xhs_output(raw: str) -> list:
    """解析小红书 MCP 搜索结果 (JSON: feeds[].noteCard.displayTitle + id)。"""
    results = []
    # bridge 返回 JSON, 但可能被截断; 先试完整解析, 失败则用正则提取
    try:
        data = json.loads(raw)
        for f in data.get("feeds", [])[:10]:
            nc = f.get("noteCard", {})
            title = nc.get("displayTitle", "") or nc.get("title", "")
            note_id = f.get("id", "")
            desc = nc.get("desc", "") or title
            user = nc.get("user", {}).get("nickname", "")
            results.append({
                "title": title[:100],
                "url": f"https://www.xiaohongshu.com/explore/{note_id}" if note_id else "",
                "snippet": (f"@{user}: {desc}" if user else desc)[:200],
            })
        if results:
            return results
    except (json.JSONDecodeError, TypeError):
        pass
    # JSON 截断时用正则兜底
    titles = re.findall(r'"displayTitle"\s*:\s*"([^"]+)"', raw)
    # note id: 24位 hex, 通常 6 开头 (小红书雪花ID)
    ids = re.findall(r'"id"\s*:\s*"(6[a-f0-9]{22,23})"', raw)
    nicknames = re.findall(r'"nickname"\s*:\s*"([^"]+)"', raw)
    for i, t in enumerate(titles[:10]):
        nid = ids[i] if i < len(ids) else ""
        nick = nicknames[i] if i < len(nicknames) else ""
        results.append({
            "title": t[:100],
            "url": f"https://www.xiaohongshu.com/explore/{nid}" if nid else "",
            "snippet": (f"@{nick}: {t}" if nick else t)[:200],
        })
    return results


def visit(url: str, goal: str = "") -> dict:
    """抓取网页正文。返回 {"url", "content", "goal"}.

    策略 (agent-reach 集成):
    1. JS 重度站点 (twitter/reddit/medium/zhihu 等) → 直接走 Jina Reader
    2. 普通站点 → requests + trafilatura
    3. trafilatura 提取失败/内容太短 → fallback Jina Reader
    """
    domain = urlparse(url).netloc.lower()

    # JS 重度站点: 直接 Jina Reader (trafilatura 抓不到 JS 渲染内容)
    if any(d in domain for d in JS_HEAVY_DOMAINS):
        return _visit_jina(url, goal)

    # 普通站点: 先试 requests + trafilatura
    try:
        headers = {"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                                 "(KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36"}
        resp = requests.get(url, headers=headers, timeout=15, allow_redirects=True,
                            proxies=PROXIES)
        resp.raise_for_status()
        if resp.encoding == "ISO-8859-1":
            resp.encoding = resp.apparent_encoding
        html = resp.text
    except Exception as e:
        # requests 失败 → 试 Jina Reader (可能是国外站点需代理, Jina 自带代理)
        jina_result = _visit_jina(url, goal)
        if jina_result.get("content"):
            return jina_result
        return {"url": url, "content": "", "error": str(e)}

    # 提取正文
    text = _html_to_text(html)
    if not text or len(text) < 50:
        # trafilatura 失败 → fallback Jina Reader
        jina_result = _visit_jina(url, goal)
        if jina_result.get("content"):
            return jina_result
        return {"url": url, "content": "", "error": "No readable content"}

    # 噪音检测: 导航菜单重复 (同一短语出现≥4次 = 噪音页面)
    # 如果噪音太大, trafilatura 提取失败, 走 Jina Reader
    if _is_navigation_noise(text):
        jina_result = _visit_jina(url, goal)
        if jina_result.get("content") and not _is_navigation_noise(jina_result["content"]):
            return jina_result

    # 截断防超长
    text = text[:8000]
    return {"url": url, "content": text, "goal": goal}


def _is_navigation_noise(text: str) -> bool:
    """检测是否为导航菜单噪音 (而非正文)。

    判断: 找出重复≥4次的 10-30 字短语, 如果重复内容占文本≥40% 就是噪音。
    """
    if len(text) < 100:
        return True
    # 提取所有 15 字片段, 统计重复
    chunks = [text[i:i+15] for i in range(0, len(text)-15, 5)]
    if not chunks:
        return False
    from collections import Counter
    common = Counter(chunks).most_common(5)
    repeat_ratio = sum(count for _, count in common if count >= 4) / len(chunks)
    return repeat_ratio > 0.3


def _visit_jina(url: str, goal: str = "") -> dict:
    """通过 Jina Reader (agent-reach web 后端) 读取网页。
    Jina Reader 能读 JS 渲染页面, 返回干净 markdown。走 mihomo 代理。
    """
    try:
        jina_url = f"{JINA_READER_URL}/{url}"
        resp = requests.get(jina_url, timeout=20, proxies=PROXIES,
                            headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        text = resp.text.strip()
        if not text or len(text) < 50:
            return {"url": url, "content": "", "error": "Jina Reader returned empty"}
        # 去掉 Jina 加的元信息头 (Title:/URL Source:/Published Time: 等)
        text = re.sub(r"^(Title|URL Source|Published Time|Warning):.*\n", "",
                      text, flags=re.MULTILINE).strip()
        return {"url": url, "content": text[:8000], "goal": goal,
                "source": "jina-reader"}
    except Exception as e:
        return {"url": url, "content": "", "error": f"Jina Reader failed: {e}"}


def _html_to_text(html: str) -> str:
    """trafilatura 提取正文, 失败则正则去标签。"""
    try:
        import trafilatura
        text = trafilatura.extract(html, include_links=True, include_tables=True,
                                   favor_recall=True)
        if text and len(text) > 100:
            return text
    except Exception:
        pass
    # 退化
    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


# ============================================================================
# 附件预处理 (MinerU PDF / DOCX / TXT → markdown)
# ============================================================================

def preprocess_attachment(filename: str, content: bytes) -> dict:
    """预处理用户上传的文档，返回 markdown 文本。

    PDF → MinerU GPU 解析 (保留表格/公式/布局)
    DOCX → python-docx 提取段落+表格
    TXT/MD → 直接读 UTF-8
    其他 → 标记不支持

    返回 {"filename", "md", "source_type", "error": optional}
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _parse_pdf_with_mineru(filename, content)
    elif ext == "docx":
        return _parse_docx(filename, content)
    elif ext in ("txt", "md", "markdown", "text"):
        return _parse_text(filename, content)
    else:
        return {"filename": filename, "md": "", "source_type": "unknown",
                "error": f"不支持的格式 .{ext}，请上传 PDF/DOCX/TXT/MD"}


def _parse_pdf_with_mineru(filename: str, content: bytes) -> dict:
    """走 MinerU GPU 解析 PDF → markdown。"""
    try:
        # 先检查 MinerU 是否可用
        health = requests.get(f"{MINERU_URL}/health", timeout=5)
        if health.status_code != 200:
            return {"filename": filename, "md": "", "source_type": "pdf",
                    "error": "MinerU 服务不可用，无法解析 PDF"}
    except Exception:
        return {"filename": filename, "md": "", "source_type": "pdf",
                "error": "MinerU 服务不可用"}

    try:
        # POST /file_parse (multipart)
        resp = requests.post(
            f"{MINERU_URL}/file_parse",
            files={"file": (filename, content, "application/pdf")},
            timeout=300,  # PDF 解析可能较慢
        )
        resp.raise_for_status()
        data = resp.json()

        # MinerU 返回格式: {"md": "..."} 或 {"pages": [{"md": "..."}]}
        if "md" in data:
            md = data["md"]
        elif "pages" in data:
            md = "\n\n".join(p.get("md", "") for p in data["pages"])
        else:
            md = json.dumps(data, ensure_ascii=False)[:2000]

        return {"filename": filename, "md": md, "source_type": "pdf"}
    except Exception as e:
        return {"filename": filename, "md": "", "source_type": "pdf",
                "error": f"MinerU 解析失败: {e}"}


def _parse_docx(filename: str, content: bytes) -> dict:
    """用 python-docx 提取段落+表格 → markdown。"""
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 简单标题检测
                if para.style and "Heading" in (para.style.name or ""):
                    level = 1
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        pass
                    lines.append(f"{'#' * level} {text}")
                else:
                    lines.append(text)
        # 表格
        for table in doc.tables:
            lines.append("\n| " + " | ".join(
                cell.text.strip() for cell in row.cells
            ) + " |" for row in table.rows)
        md = "\n\n".join(lines)
        return {"filename": filename, "md": md, "source_type": "docx"}
    except ImportError:
        return {"filename": filename, "md": "", "source_type": "docx",
                "error": "python-docx 未安装"}
    except Exception as e:
        return {"filename": filename, "md": "", "source_type": "docx",
                "error": f"DOCX 解析失败: {e}"}


def _parse_text(filename: str, content: bytes) -> dict:
    """直接读文本文件。"""
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            text = content.decode(encoding)
            return {"filename": filename, "md": text, "source_type": "text"}
        except (UnicodeDecodeError, ValueError):
            continue
    return {"filename": filename, "md": "", "source_type": "text",
            "error": "无法解码文件"}


# ============================================================================
# 幻觉抵御: 事实声明验证
# ============================================================================

def verify_claim(claim: str) -> dict:
    """验证一个事实声明: 构造搜索 query → SearXNG 搜索 → 检查多来源支持。

    返回 {"status": "verified/partial/unverified", "evidence": [...]}
    """
    # 构造搜索 query: 去停用词 + 保留数字/专有名词
    query = _claim_to_query(claim)
    if not query:
        return {"status": "unverified", "evidence": [], "reason": "无法构造验证查询"}

    # 搜索
    try:
        lang = "zh" if any("\u4E00" <= c <= "\u9FFF" for c in query) else "en"
        resp = requests.get(
            f"{SEARXNG_URL}/search",
            params={"q": query, "format": "json", "categories": "general",
                    "pageno": 1, "language": lang},
            timeout=15,
        )
        resp.raise_for_status()
        results = resp.json().get("results", [])
    except Exception:
        return {"status": "unverified", "evidence": [], "reason": "搜索失败"}

    # 检查结果是否支持该声明 (关键词匹配)
    claim_keywords = _extract_keywords(claim)
    supporting = []
    for r in results[:5]:
        snippet = (r.get("title", "") + " " + r.get("content", "")).lower()
        # 如果 snippet 包含声明中 ≥50% 的关键词, 认为支持
        matched = sum(1 for kw in claim_keywords if kw.lower() in snippet)
        if claim_keywords and matched >= max(2, len(claim_keywords) * 0.5):
            supporting.append({
                "title": r.get("title", ""),
                "url": r.get("url", ""),
                "snippet": r.get("content", "")[:150],
            })

    if len(supporting) >= 2:
        return {"status": "verified", "evidence": supporting[:3]}
    elif len(supporting) == 1:
        return {"status": "partial", "evidence": supporting}
    else:
        return {"status": "unverified", "evidence": [],
                "reason": "未找到支持该声明的独立来源"}


def verify_url(url: str) -> dict:
    """检查引用 URL 是否可达且有内容。走代理验证国外链接。"""
    try:
        resp = requests.head(url, timeout=8, allow_redirects=True,
                             headers={"User-Agent": "Mozilla/5.0"}, proxies=PROXIES)
        if resp.status_code < 400:
            return {"reachable": True, "status_code": resp.status_code}
        return {"reachable": False, "status_code": resp.status_code}
    except Exception:
        return {"reachable": False, "status_code": 0}


def _claim_to_query(claim: str) -> str:
    """把事实声明转成搜索 query: 去停用词 + 保留关键信息。"""
    # 提取带引号的短语 + 数字 + 大写词
    keywords = _extract_keywords(claim)
    if not keywords:
        return claim[:60]
    return " ".join(keywords[:6])


def _extract_keywords(text: str) -> list:
    """提取关键词: 去停用词, 保留数字/专有名词/中文实词。"""
    # 英文: 按空格分词, 去停用词, 保留含数字或首字母大写的词
    words = re.findall(r"\b[A-Za-z]+\d*|\d+(?:\.\d+)?%?|[\u4e00-\u9fff]+", text)
    keywords = []
    for w in words:
        wl = w.lower()
        if wl in STOP_WORDS or len(wl) < 2:
            continue
        # 保留: 含数字 / 首字母大写 / 中文词 (≥2字)
        if any(c.isdigit() for c in w) or (w[0].isupper() and len(w) > 1) or any("\u4e00" <= c <= "\u9fff" for c in w):
            keywords.append(w)
    return keywords
