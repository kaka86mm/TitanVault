"""
server.py — Deep Research FastAPI 服务

端点:
  POST   /api/sessions               新建研究 session
  GET    /api/sessions               session 列表
  GET    /api/sessions/{id}          session 详情
  POST   /api/sessions/{id}/message  迭代: 发送 follow-up
  GET    /api/sessions/{id}/stream   SSE 实时进度流
  POST   /api/sessions/{id}/finalize 标记完成
  DELETE /api/sessions/{id}          删除 session
  POST   /api/ingest                 存入 Open Notebook
  GET    /api/health                 健康检查
  GET    /                            前端页面
"""
import os
import json
import asyncio
import threading
import uuid as uuid_lib
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from agent import ResearchAgent, ResearchContext
from tools import preprocess_attachment
import scheduler as sched
import base64

# ============================================================================
# 配置
# ============================================================================

QUEST_ENDPOINT = os.environ.get("QUEST_ENDPOINT", "http://host-gateway:8093/v1")
QUEST_MODEL = os.environ.get("QUEST_MODEL", "QUEST-9B")
LITELLM_URL = os.environ.get("LITELLM_URL", "http://host-gateway:4000/v1")
LITELLM_KEY = os.environ.get("LITELLM_KEY", os.environ.get("LITELLM_MASTER_KEY", "EMPTY"))
ON_BASE_URL = os.environ.get("ON_BASE_URL", "http://host-gateway:5055")
REPORTS_DIR = Path(os.environ.get("REPORTS_DIR", "/data/quest-reports"))
SESSIONS_DIR = REPORTS_DIR / ".sessions"
MAX_TURNS = int(os.environ.get("MAX_TURNS", "5"))

REPORTS_DIR.mkdir(parents=True, exist_ok=True)
SESSIONS_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="TitanVault Deep Research", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_methods=["*"], allow_headers=["*"])

# ============================================================================
# Session 存储 (内存 + 磁盘持久化)
# ============================================================================

sessions: Dict[str, ResearchContext] = {}
# 每个 session 的 SSE 事件队列 (运行中的任务推送事件)
event_queues: Dict[str, asyncio.Queue] = {}
# 运行状态
session_status: Dict[str, dict] = {}


def _save_session(ctx: ResearchContext):
    """持久化 session 到磁盘。"""
    p = SESSIONS_DIR / f"{ctx.session_id}.json"
    p.write_text(json.dumps(ctx.to_dict(), ensure_ascii=False, indent=2))


def _load_session(sid: str) -> Optional[ResearchContext]:
    """从磁盘加载 session。"""
    p = SESSIONS_DIR / f"{sid}.json"
    if p.exists():
        return ResearchContext.from_dict(json.loads(p.read_text()))
    return None


def _load_all_sessions():
    """启动时加载所有持久化的 session。"""
    if not SESSIONS_DIR.exists():
        return
    for p in SESSIONS_DIR.glob("*.json"):
        try:
            ctx = ResearchContext.from_dict(json.loads(p.read_text()))
            sessions[ctx.session_id] = ctx
        except Exception:
            pass


_load_all_sessions()


# ============================================================================
# 定时研究调度器
# ============================================================================

def _scheduled_run_callback(sid: str, question: str, attachments=None,
                             is_scheduled=False):
    """定时任务回调: 创建 session + 跑 agent (同步, 在 scheduler 线程)。"""
    ctx = ResearchContext(sid)
    if attachments:
        for att in attachments:
            ctx.add_attachment(att.get("filename", ""), att.get("md", ""),
                              att.get("source_type", "text"))
    sessions[sid] = ctx
    event_queues[sid] = asyncio.Queue()
    _save_session(ctx)
    _run_agent_thread(sid, question, False)


sched.init_scheduler(_scheduled_run_callback)
# ============================================================================

class CreateSession(BaseModel):
    question: str
    attachments: Optional[list] = None  # [{filename, content_b64}]


class SendMessage(BaseModel):
    content: str
    attachments: Optional[list] = None


class IngestRequest(BaseModel):
    session_id: str
    version: Optional[int] = None  # 默认最新版本


# ============================================================================
# Agent 运行 (后台线程)
# ============================================================================

def _run_agent_thread(session_id: str, question: str, is_followup: bool = False):
    """在后台线程跑 agent, 事件推入 asyncio.Queue。"""
    ctx = sessions[session_id]
    loop = asyncio.new_event_loop()

    # 确保 event_queue 存在 (定时任务可能没预创建)
    if session_id not in event_queues:
        event_queues[session_id] = asyncio.Queue()

    async def push(event):
        q = event_queues.get(session_id)
        if q:
            # 从线程安全地推入主线程的 queue
            try:
                event_queues[session_id].put_nowait(event)
            except Exception:
                pass

    def on_event(event):
        # agent 回调 (同步线程里) → 写入 queue
        # asyncio.Queue 跨线程用 put_nowait (线程安全)
        q = event_queues.get(session_id)
        if q:
            try:
                q.put_nowait(event)
            except Exception:
                pass

    agent = ResearchAgent(
        endpoint=QUEST_ENDPOINT, model=QUEST_MODEL, max_turns=MAX_TURNS,
    )

    session_status[session_id] = {"status": "running", "turn": 0}

    try:
        ctx.add_message(question, is_followup=is_followup)
        report = agent.run(question, context=ctx, on_event=on_event)
        # 保存报告到文件
        safe_q = "".join(c if c.isalnum() or c in "-_" else "_" for c in question[:30])
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        report_file = REPORTS_DIR / f"rs-{session_id}-v{ctx.current_version}-{ts}.md"
        report_file.write_text(report)
        session_status[session_id] = {"status": "done", "turn": agent.max_turns}
        on_event({"type": "done", "report_file": str(report_file),
                  "version": ctx.current_version})
    except Exception as e:
        session_status[session_id] = {"status": "failed", "error": str(e)}
        on_event({"type": "error", "message": str(e)})
    finally:
        _save_session(ctx)
        on_event({"type": "end"})


# ============================================================================
# Routes
# ============================================================================

@app.get("/api/health")
async def health():
    return {"status": "ok", "model": QUEST_MODEL, "sessions": len(sessions)}


@app.post("/api/sessions")
async def create_session(req: CreateSession):
    sid = f"rs_{uuid_lib.uuid4().hex[:12]}"
    ctx = ResearchContext(sid)
    sessions[sid] = ctx
    event_queues[sid] = asyncio.Queue()

    # 处理附件 (预处理: PDF→MinerU, DOCX→python-docx, TXT→直读)
    if req.attachments:
        for att in req.attachments:
            filename = att.get("filename", "unknown")
            content_b64 = att.get("content_b64", "")
            try:
                content_bytes = base64.b64decode(content_b64)
                result = preprocess_attachment(filename, content_bytes)
                if result.get("md"):
                    ctx.add_attachment(filename, result["md"],
                                      result.get("source_type", "text"))
                else:
                    ctx.add_attachment(filename,
                                      f"[附件解析失败: {result.get('error','')}]", "error")
            except Exception as e:
                ctx.add_attachment(filename, f"[附件处理异常: {e}]", "error")

    _save_session(ctx)

    # 后台启动 agent
    t = threading.Thread(target=_run_agent_thread, args=(sid, req.question, False),
                         daemon=True)
    t.start()

    return {"session_id": sid, "status": "running", "question": req.question,
            "attachments": len(ctx.attachments)}


@app.get("/api/sessions")
async def list_sessions():
    result = []
    for sid, ctx in sessions.items():
        result.append({
            "session_id": sid,
            "created_at": ctx.created_at,
            "status": ctx.status,
            "current_version": ctx.current_version,
            "question": ctx.messages[0]["content"] if ctx.messages else "",
            "n_messages": len(ctx.messages),
        })
    result.sort(key=lambda x: x["created_at"], reverse=True)
    return result


@app.get("/api/sessions/{sid}")
async def get_session(sid: str):
    ctx = sessions.get(sid)
    if not ctx:
        ctx = _load_session(sid)
        if ctx:
            sessions[sid] = ctx
        else:
            raise HTTPException(404, "Session not found")
    status = session_status.get(sid, {"status": "idle"})
    return {**ctx.to_dict(), "run_status": status}


@app.post("/api/sessions/{sid}/message")
async def send_message(sid: str, req: SendMessage):
    ctx = sessions.get(sid)
    if not ctx:
        raise HTTPException(404, "Session not found")
    if ctx.status == "finalized":
        raise HTTPException(400, "Session is finalized")
    if session_status.get(sid, {}).get("status") == "running":
        raise HTTPException(409, "Session is already running")

    # 处理迭代时用户补充的附件 (预处理: PDF→MinerU, DOCX→python-docx, TXT→直读)
    if req.attachments:
        for att in req.attachments:
            filename = att.get("filename", "unknown")
            content_b64 = att.get("content_b64", "")
            try:
                content_bytes = base64.b64decode(content_b64)
                result = preprocess_attachment(filename, content_bytes)
                if result.get("md"):
                    ctx.add_attachment(filename, result["md"],
                                      result.get("source_type", "text"))
                else:
                    ctx.add_attachment(filename,
                                      f"[附件解析失败: {result.get('error','')}]", "error")
            except Exception as e:
                ctx.add_attachment(filename, f"[附件处理异常: {e}]", "error")

    # 重置事件队列
    event_queues[sid] = asyncio.Queue()

    # 后台启动迭代
    t = threading.Thread(target=_run_agent_thread, args=(sid, req.content, True),
                         daemon=True)
    t.start()

    return {"session_id": sid, "status": "running", "message": req.content,
            "attachments": len(ctx.attachments)}


@app.get("/api/sessions/{sid}/stream")
async def stream_session(sid: str):
    """SSE 流: 实时推送 agent 事件。"""
    if sid not in event_queues:
        raise HTTPException(404, "No active stream for this session")

    async def event_generator():
        q = event_queues[sid]
        while True:
            try:
                event = await asyncio.wait_for(q.get(), timeout=120)
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
                if event.get("type") in ("end", "error"):
                    break
            except asyncio.TimeoutError:
                yield f"data: {json.dumps({'type': 'heartbeat'})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache",
                                      "X-Accel-Buffering": "no"})


@app.post("/api/sessions/{sid}/finalize")
async def finalize_session(sid: str):
    ctx = sessions.get(sid)
    if not ctx:
        raise HTTPException(404, "Session not found")
    ctx.status = "finalized"
    _save_session(ctx)
    return {"session_id": sid, "status": "finalized"}


@app.delete("/api/sessions/{sid}")
async def delete_session(sid: str):
    if sid in sessions:
        del sessions[sid]
    p = SESSIONS_DIR / f"{sid}.json"
    if p.exists():
        p.unlink()
    return {"deleted": sid}


@app.post("/api/ingest")
async def ingest_report(req: IngestRequest):
    """存入 Open Notebook (串联 ingest 能力)。"""
    ctx = sessions.get(req.session_id)
    if not ctx:
        raise HTTPException(404, "Session not found")

    version = req.version or ctx.current_version
    version_data = next((v for v in ctx.versions if v["version"] == version), None)
    if not version_data:
        raise HTTPException(404, f"Version {version} not found")

    content = version_data["content"]
    title = ctx.messages[0]["content"][:50] if ctx.messages else "Research Report"

    # POST 到 Open Notebook
    import httpx
    async with httpx.AsyncClient() as client:
        try:
            base = ON_BASE_URL.rstrip("/")
            # 探测路径前缀
            for prefix in ["/api", ""]:
                r = await client.post(
                    f"{base}{prefix}/sources",
                    data={"type": "text", "content": content, "embed": "true",
                          "async_processing": "true", "title": f"研究: {title}"},
                    timeout=30,
                )
                if r.status_code in (200, 201):
                    return {"status": "ok", "source_id": r.json().get("id")}
            return {"status": "error", "code": r.status_code, "body": r.text[:200]}
        except Exception as e:
            return {"status": "error", "message": str(e)}


def _cd_header(title: str, ext: str) -> dict:
    """生成 Content-Disposition header (支持中文文件名)。"""
    from urllib.parse import quote
    safe = quote(f"{title}.{ext}")
    return {"Content-Disposition": f"attachment; filename*=UTF-8''{safe}"}


# ============================================================================
# 报告导出
# ============================================================================

@app.get("/api/sessions/{sid}/export")
async def export_report(sid: str, format: str = "md"):
    """导出研究报告为 Markdown / PDF / DOCX。"""
    ctx = sessions.get(sid)
    if not ctx:
        raise HTTPException(404, "Session not found")
    if not ctx.versions:
        raise HTTPException(404, "No report available")

    report = ctx.versions[-1].get("content", "")
    title = ctx.messages[0]["content"][:50] if ctx.messages else "Research Report"

    if format == "md":
        from fastapi.responses import Response
        return Response(content=report.encode("utf-8"),
                        media_type="text/markdown",
                        headers=_cd_header(title, "md"))

    elif format == "pdf":
        return _export_pdf(report, title)

    elif format == "docx":
        return _export_docx(report, title)

    raise HTTPException(400, f"Unsupported format: {format}")


def _export_pdf(report: str, title: str):
    """用 reportlab 生成 PDF (内置 CID 中文字体 STSong-Light)。"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase import pdfmetrics
        from io import BytesIO
        import re as _re
        from xml.sax.saxutils import escape as _xml_escape

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)
        h1 = ParagraphStyle("H1", fontName="STSong-Light", fontSize=18, leading=26, spaceAfter=12, spaceBefore=18)
        h2 = ParagraphStyle("H2", fontName="STSong-Light", fontSize=14, leading=20, spaceAfter=8, spaceBefore=16)
        h3 = ParagraphStyle("H3", fontName="STSong-Light", fontSize=12, leading=18, spaceAfter=6, spaceBefore=12)
        body = ParagraphStyle("Body", fontName="STSong-Light", fontSize=10.5, leading=17, spaceAfter=4)
        quote = ParagraphStyle("Quote", fontName="STSong-Light", fontSize=9.5, leading=15, leftIndent=15, textColor="#666666")

        story = []
        for line in report.split("\n"):
            line = line.rstrip()
            if not line:
                story.append(Spacer(1, 4))
                continue
            is_h1 = line.startswith("# ")
            is_h2 = line.startswith("## ")
            is_h3 = line.startswith("### ")
            is_quote = line.startswith("> ")
            # 清理 markdown → reportlab XML
            clean = _re.sub(r'\[([^\]]*)\]\(([^\)]+)\)', r'<a href="\2" color="#3b9eff">\1</a>', line)
            clean = _re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', clean)
            clean = _re.sub(r'^#+\s*', '', clean)
            clean = _re.sub(r'^>\s*', '', clean)
            clean = _re.sub(r'^[-*]\s+', '- ', clean)
            # escape & < > 但保留已生成的 XML 标签
            clean = _re.sub(r'&(?!amp;|lt;|gt;|quot;|#)', '&amp;', clean)
            if not clean.strip():
                continue
            try:
                if is_h1: story.append(Paragraph(clean, h1))
                elif is_h2: story.append(Paragraph(clean, h2))
                elif is_h3: story.append(Paragraph(clean, h3))
                elif is_quote: story.append(Paragraph(clean, quote))
                else: story.append(Paragraph(clean, body))
            except Exception:
                try: story.append(Paragraph(_xml_escape(clean), body))
                except Exception: continue

        doc.build(story)
        buf.seek(0)
        from fastapi.responses import StreamingResponse
        return StreamingResponse(buf, media_type="application/pdf",
                                 headers=_cd_header(title, "pdf"))
    except Exception as e:
        import logging
        logging.error(f"PDF export failed: {e}", exc_info=True)
        md_html = _md_to_html(report)
        full_html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{md_html}</body></html>"
        from fastapi.responses import Response
        return Response(content=full_html.encode("utf-8"),
                        media_type="text/html",
                        headers=_cd_header(title, "html"))


def _export_docx(report: str, title: str):
    """用 python-docx 生成 Word 文档。"""
    try:
        from docx import Document
        from io import BytesIO
        import re as _re

        doc = Document()
        doc.core_properties.title = title
        _table_ref = [None]

        for line in report.split("\n"):
            line = line.rstrip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.strip().startswith("|") and "|" in line.strip()[1:]:
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                if all(_re.match(r'^[-:]+$', c) for c in cells):
                    continue
                if _table_ref[0] is None:
                    _table_ref[0] = doc.add_table(rows=0, cols=len(cells))
                    _table_ref[0].style = 'Table Grid'
                row = _table_ref[0].add_row()
                for i, cell in enumerate(cells):
                    if i < len(row.cells):
                        clean = _re.sub(r'\[([^\]]*)\]\([^\)]+\)', r'\1', cell)
                        row.cells[i].text = clean
            elif line.startswith("> "):
                doc.add_paragraph(line[2:].strip())
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            else:
                clean = _re.sub(r'\[([^\]]*)\]\([^\)]+\)', r'\1', line)
                clean = _re.sub(r'\*\*([^*]+)\*\*', r'\1', clean)
                clean = _re.sub(r'`([^`]+)`', r'\1', clean)
                doc.add_paragraph(clean)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        from fastapi.responses import StreamingResponse
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=_cd_header(title, "docx"))
    except ImportError:
        raise HTTPException(500, "python-docx not installed")


def _md_to_html(md: str) -> str:
    """简单 markdown → HTML。"""
    import html as _html
    import re as _re
    lines = md.split("\n")
    out = []
    in_table = False
    in_code = False
    for line in lines:
        if line.strip().startswith("```"):
            out.append("</code></pre>" if in_code else "<pre><code>")
            in_code = not in_code
            continue
        if in_code:
            out.append(_html.escape(line))
            continue
        line = _html.escape(line)
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(_re.match(r'^[-:]+$', c) for c in cells):
                continue
            tag = "th" if not in_table else "td"
            if not in_table:
                out.append("<table>"); in_table = True
            out.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
            continue
        elif in_table:
            out.append("</table>"); in_table = False
        if line.startswith("### "): out.append(f"<h3>{line[4:]}</h3>")
        elif line.startswith("## "): out.append(f"<h2>{line[3:]}</h2>")
        elif line.startswith("# "): out.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith("> "): out.append(f"<blockquote>{line[2:]}</blockquote>")
        elif line.startswith("- ") or line.startswith("* "): out.append(f"<li>{line[2:]}</li>")
        elif line.strip() == "---": out.append("<hr>")
        elif line.strip():
            line = _re.sub(r'\[([^\]]*)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', line)
            line = _re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', line)
            line = _re.sub(r'`([^`]+)`', r'<code>\1</code>', line)
            out.append(f"<p>{line}</p>")
    if in_table: out.append("</table>")
    return "\n".join(out)


# ============================================================================
# 定时研究
# ============================================================================

class CreateSchedule(BaseModel):
    question: str
    interval: str = "daily"  # daily/weekly/12h/6h/数字
    focus: str = ""
    attachments: Optional[list] = None


@app.post("/api/schedules")
async def create_schedule(req: CreateSchedule):
    # 预处理附件
    processed_attachments = []
    if req.attachments:
        for att in req.attachments:
            filename = att.get("filename", "unknown")
            content_b64 = att.get("content_b64", "")
            try:
                content_bytes = base64.b64decode(content_b64)
                result = preprocess_attachment(filename, content_bytes)
                if result.get("md"):
                    processed_attachments.append({
                        "filename": filename,
                        "md": result["md"][:5000],
                        "source_type": result.get("source_type", "text"),
                    })
            except Exception:
                pass

    meta = sched.add_schedule(
        question=req.question,
        interval=req.interval,
        focus=req.focus,
        attachments=processed_attachments,
    )
    return meta


@app.get("/api/schedules")
async def list_schedules():
    return sched.list_schedules()


@app.delete("/api/schedules/{schedule_id}")
async def delete_schedule(schedule_id: str):
    ok = sched.remove_schedule(schedule_id)
    if not ok:
        raise HTTPException(404, "Schedule not found")
    return {"deleted": schedule_id}


@app.post("/api/schedules/{schedule_id}/pause")
async def pause_schedule(schedule_id: str):
    ok = sched.pause_schedule(schedule_id)
    if not ok:
        raise HTTPException(404, "Schedule not found")
    return {"paused": schedule_id}


@app.post("/api/schedules/{schedule_id}/resume")
async def resume_schedule(schedule_id: str):
    ok = sched.resume_schedule(schedule_id)
    if not ok:
        raise HTTPException(404, "Schedule not found")
    return {"resumed": schedule_id}


@app.get("/", response_class=HTMLResponse)
async def index():
    p = Path(__file__).parent / "index.html"
    return HTMLResponse(p.read_text())


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", "8099")))
