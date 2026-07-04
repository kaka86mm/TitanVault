"""
tools.py — Deep Research 工具集 (SearXNG 搜索 + 网页阅读)

本地化实现, 无外部付费 API。
- search: SearXNG 元搜索引擎
- visit: requests + trafilatura 网页抓取
"""
import os
import re
import json
import requests
from typing import List, Optional
from urllib.parse import urlparse

SEARXNG_URL = os.environ.get("SEARXNG_URL", "http://host-gateway:8087")
MINERU_URL = os.environ.get("MINERU_URL", "http://host-gateway:18080")
# 已知需要 JS 渲染的域名 (走 Chrome CDP, 当前简化为标记跳过)
JS_HEAVY_DOMAINS = {"twitter.com", "x.com", "reddit.com", "instagram.com"}

# 验证用的停用词 (构造 verify query 时去掉)
STOP_WORDS = frozenset(
    "the a an is are was were be been being have has had do does did "
    "will would could should may might must can of to in on at for with "
    "and or not but if then that this these those it its as from by "
    "的 了 是 在 和 与 或 也 就 都 还 不 没 有 对 从 被 把 让 "
    "what which who when where why how about into than".split()
)


def search(query: str, skip_queries: list = None) -> dict:
    """SearXNG 搜索。返回 {"query", "results": [{"title","url","snippet"}]}.

    skip_queries: 已搜过的 query 列表 (去重, 避免迭代时重复搜索)。
    """
    if skip_queries and query.lower() in [q.lower() for q in skip_queries]:
        return {"query": query, "results": [], "skipped": True}

    lang = "zh" if any("\u4E00" <= c <= "\u9FFF" for c in query) else "en"
    try:
        resp = requests.get(
            f"{SEARXNG_URL}/search",
            params={"q": query, "format": "json", "categories": "general",
                    "pageno": 1, "language": lang},
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        return {"query": query, "results": [], "error": str(e)}

    results = []
    for r in data.get("results", [])[:10]:
        results.append({
            "title": r.get("title", ""),
            "url": r.get("url", ""),
            "snippet": r.get("content", "")[:200],
        })
    return {"query": query, "results": results}


def visit(url: str, goal: str = "") -> dict:
    """抓取网页正文。返回 {"url", "content", "goal"}."""
    # 检查 JS 重度站点
    domain = urlparse(url).netloc.lower()
    if any(d in domain for d in JS_HEAVY_DOMAINS):
        return {"url": url, "content": "", "error": "JS-heavy site, skipped"}

    # 抓取 HTML
    try:
        headers = {"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                                 "(KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36"}
        resp = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        resp.raise_for_status()
        if resp.encoding == "ISO-8859-1":
            resp.encoding = resp.apparent_encoding
        html = resp.text
    except Exception as e:
        return {"url": url, "content": "", "error": str(e)}

    # 提取正文
    text = _html_to_text(html)
    if not text or len(text) < 50:
        return {"url": url, "content": "", "error": "No readable content"}

    # 截断防超长
    text = text[:8000]
    return {"url": url, "content": text, "goal": goal}


def _html_to_text(html: str) -> str:
    """trafilatura 提取正文, 失败则正则去标签。"""
    try:
        import trafilatura
        text = trafilatura.extract(html, include_links=True, include_tables=True,
                                   favor_recall=True)
        if text and len(text) > 100:
            return text
    except Exception:
        pass
    # 退化
    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


# ============================================================================
# 附件预处理 (MinerU PDF / DOCX / TXT → markdown)
# ============================================================================

def preprocess_attachment(filename: str, content: bytes) -> dict:
    """预处理用户上传的文档，返回 markdown 文本。

    PDF → MinerU GPU 解析 (保留表格/公式/布局)
    DOCX → python-docx 提取段落+表格
    TXT/MD → 直接读 UTF-8
    其他 → 标记不支持

    返回 {"filename", "md", "source_type", "error": optional}
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _parse_pdf_with_mineru(filename, content)
    elif ext == "docx":
        return _parse_docx(filename, content)
    elif ext in ("txt", "md", "markdown", "text"):
        return _parse_text(filename, content)
    else:
        return {"filename": filename, "md": "", "source_type": "unknown",
                "error": f"不支持的格式 .{ext}，请上传 PDF/DOCX/TXT/MD"}


def _parse_pdf_with_mineru(filename: str, content: bytes) -> dict:
    """走 MinerU GPU 解析 PDF → markdown。"""
    try:
        # 先检查 MinerU 是否可用
        health = requests.get(f"{MINERU_URL}/health", timeout=5)
        if health.status_code != 200:
            return {"filename": filename, "md": "", "source_type": "pdf",
                    "error": "MinerU 服务不可用，无法解析 PDF"}
    except Exception:
        return {"filename": filename, "md": "", "source_type": "pdf",
                "error": "MinerU 服务不可用"}

    try:
        # POST /file_parse (multipart)
        resp = requests.post(
            f"{MINERU_URL}/file_parse",
            files={"file": (filename, content, "application/pdf")},
            timeout=300,  # PDF 解析可能较慢
        )
        resp.raise_for_status()
        data = resp.json()

        # MinerU 返回格式: {"md": "..."} 或 {"pages": [{"md": "..."}]}
        if "md" in data:
            md = data["md"]
        elif "pages" in data:
            md = "\n\n".join(p.get("md", "") for p in data["pages"])
        else:
            md = json.dumps(data, ensure_ascii=False)[:2000]

        return {"filename": filename, "md": md, "source_type": "pdf"}
    except Exception as e:
        return {"filename": filename, "md": "", "source_type": "pdf",
                "error": f"MinerU 解析失败: {e}"}


def _parse_docx(filename: str, content: bytes) -> dict:
    """用 python-docx 提取段落+表格 → markdown。"""
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 简单标题检测
                if para.style and "Heading" in (para.style.name or ""):
                    level = 1
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        pass
                    lines.append(f"{'#' * level} {text}")
                else:
                    lines.append(text)
        # 表格
        for table in doc.tables:
            lines.append("\n| " + " | ".join(
                cell.text.strip() for cell in row.cells
            ) + " |" for row in table.rows)
        md = "\n\n".join(lines)
        return {"filename": filename, "md": md, "source_type": "docx"}
    except ImportError:
        return {"filename": filename, "md": "", "source_type": "docx",
                "error": "python-docx 未安装"}
    except Exception as e:
        return {"filename": filename, "md": "", "source_type": "docx",
                "error": f"DOCX 解析失败: {e}"}


def _parse_text(filename: str, content: bytes) -> dict:
    """直接读文本文件。"""
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            text = content.decode(encoding)
            return {"filename": filename, "md": text, "source_type": "text"}
        except (UnicodeDecodeError, ValueError):
            continue
    return {"filename": filename, "md": "", "source_type": "text",
            "error": "无法解码文件"}


# ============================================================================
# 幻觉抵御: 事实声明验证
# ============================================================================

def verify_claim(claim: str) -> dict:
    """验证一个事实声明: 构造搜索 query → SearXNG 搜索 → 检查多来源支持。

    返回 {"status": "verified/partial/unverified", "evidence": [...]}
    """
    # 构造搜索 query: 去停用词 + 保留数字/专有名词
    query = _claim_to_query(claim)
    if not query:
        return {"status": "unverified", "evidence": [], "reason": "无法构造验证查询"}

    # 搜索
    try:
        lang = "zh" if any("\u4E00" <= c <= "\u9FFF" for c in query) else "en"
        resp = requests.get(
            f"{SEARXNG_URL}/search",
            params={"q": query, "format": "json", "categories": "general",
                    "pageno": 1, "language": lang},
            timeout=15,
        )
        resp.raise_for_status()
        results = resp.json().get("results", [])
    except Exception:
        return {"status": "unverified", "evidence": [], "reason": "搜索失败"}

    # 检查结果是否支持该声明 (关键词匹配)
    claim_keywords = _extract_keywords(claim)
    supporting = []
    for r in results[:5]:
        snippet = (r.get("title", "") + " " + r.get("content", "")).lower()
        # 如果 snippet 包含声明中 ≥50% 的关键词, 认为支持
        matched = sum(1 for kw in claim_keywords if kw.lower() in snippet)
        if claim_keywords and matched >= max(2, len(claim_keywords) * 0.5):
            supporting.append({
                "title": r.get("title", ""),
                "url": r.get("url", ""),
                "snippet": r.get("content", "")[:150],
            })

    if len(supporting) >= 2:
        return {"status": "verified", "evidence": supporting[:3]}
    elif len(supporting) == 1:
        return {"status": "partial", "evidence": supporting}
    else:
        return {"status": "unverified", "evidence": [],
                "reason": "未找到支持该声明的独立来源"}


def verify_url(url: str) -> dict:
    """检查引用 URL 是否可达且有内容。"""
    try:
        resp = requests.head(url, timeout=8, allow_redirects=True,
                             headers={"User-Agent": "Mozilla/5.0"})
        if resp.status_code < 400:
            return {"reachable": True, "status_code": resp.status_code}
        return {"reachable": False, "status_code": resp.status_code}
    except Exception:
        return {"reachable": False, "status_code": 0}


def _claim_to_query(claim: str) -> str:
    """把事实声明转成搜索 query: 去停用词 + 保留关键信息。"""
    # 提取带引号的短语 + 数字 + 大写词
    keywords = _extract_keywords(claim)
    if not keywords:
        return claim[:60]
    return " ".join(keywords[:6])


def _extract_keywords(text: str) -> list:
    """提取关键词: 去停用词, 保留数字/专有名词/中文实词。"""
    # 英文: 按空格分词, 去停用词, 保留含数字或首字母大写的词
    words = re.findall(r"\b[A-Za-z]+\d*|\d+(?:\.\d+)?%?|[\u4e00-\u9fff]+", text)
    keywords = []
    for w in words:
        wl = w.lower()
        if wl in STOP_WORDS or len(wl) < 2:
            continue
        # 保留: 含数字 / 首字母大写 / 中文词 (≥2字)
        if any(c.isdigit() for c in w) or (w[0].isupper() and len(w) > 1) or any("\u4e00" <= c <= "\u9fff" for c in w):
            keywords.append(w)
    return keywords
